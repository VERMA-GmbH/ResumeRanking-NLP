import os
import re
import sys
from io import StringIO
import time
import openai
import docx
from docx import Document
from PyPDF2 import PdfReader
import json




open_ai_key = os.environ['openapi_key']
openai.api_key = open_ai_key

def convert_pdf_to_docx(pdf_path, docx_path):
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PdfReader(pdf_file)
        output = StringIO()
        for i in range(len(pdf_reader.pages)):
            output.write(pdf_reader.pages[i].extract_text())
        text = output.getvalue()
        output.close()

    doc = Document()
    for line in text.split("\n"):
        if re.match(r"\s*\d+\.\d+\s+", line):
            paragraph = doc.add_paragraph(line)
            paragraph.style = doc.styles["List Bullet"]
        else:
            paragraph = doc.add_paragraph(line)

    doc.save(docx_path)




def check_pdf_file(pdf_path):
    pdf_extensions = [
        ".pdf",
        ".PDF"
    ]
    file_name, file_extension = os.path.splitext(pdf_path)

    if file_extension in pdf_extensions:
        return True, file_name + ".docx"
    return False, pdf_path


def check_and_convert_pdf_file(pdf_path, inplace = True):
    try:
        check_result, file_name = check_pdf_file(pdf_path)
        if check_result:
            convert_pdf_to_docx(pdf_path, file_name)
            if inplace:
                os.remove(pdf_path)
        return file_name
    except Exception as e: 
        os.remove(pdf_path)
        print(f"Error in converting PDF {pdf_path} to docx\n", e)
        return False
    



def delete_dir(dir_path = None, delete_before = 24*60*60):
    # Specify the directory path
    if dir_path == None:
        return
    # Get the current time in seconds
    now = time.time()

    # Loop through all directories in the specified path
    for dir in os.listdir(dir_path):
        # Get the full path of the directory
        dir_full_path = os.path.join(dir_path, dir)
        
        # Check if the directory exists and is actually a directory
        if os.path.isdir(dir_full_path):
            # Get the creation time of the directory in seconds
            creation_time = os.path.getctime(dir_full_path)
            
            # Calculate the age of the directory in seconds
            age = now - creation_time
            
            # Check if the directory is older than 24 hours
            if age > delete_before:
                # Delete the directory and its contents recursively
                os.system("rm -rf " + dir_full_path)


def del_old_data(path="/root/ResumeRanking-NLP/Data", delete_before = 24*60*60):
    for folder in ["JobDesc", "Resumes"]:
        delete_dir(os.path.join(path, folder), delete_before)


def read_docx_file(filename):
    """
    Reads the contents of a docx file and returns a string containing all the text in the file.
    """
    # Load the docx file
    doc = docx.Document(filename)

    # Get all the text from the file
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)

    # Return the text as a single string
    return '\n'.join(all_text)

def extract_mobile_numbers(text):
    """
    Extracts mobile numbers from a string and returns them as a list.
    """
    # Define the regular expression for a mobile number
    pattern = r'\b\d{10}\b'

    # Search for the pattern in the text
    match_iter = re.finditer(pattern, text)

    # Extract the mobile numbers from the matches
    mobile_numbers = []
    for match in match_iter:
        mobile_numbers.append(match.group())
    # Return the mobile numbers as a list
    return mobile_numbers


def extract_email_addresses(text):
    """
    Extracts email addresses from a string and returns them as a list.
    """
    # Define the regular expression pattern
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

    # Search for the pattern in the text
    matches = list(re.finditer(pattern, text))

    # Select the email address with the shorter length in case of overlapping matches
    email_addresses = []
    for i, match in enumerate(matches):
        try:
            current_address = match.group()
            if i == 0  :
                if current_address not in email_addresses:
                    email_addresses.append(current_address)
            else:
                previous_match = matches[i-1]
                previous_address = previous_match.group()
                if match.start() < previous_match.end():
                    if len(current_address) < len(previous_address):
                        email_addresses[-1] = current_address
                else:
                    if current_address not in email_addresses:
                        email_addresses.append(current_address)
        except Exception as e:
            print(e)
    filtered_email = []
    for email in email_addresses:
        if email.endswith(".com"):
            filtered_email.append(email)
    # Return the email addresses as a list
    return filtered_email

def get_similarity_post_processing(data):
    data["email"] = "N/A"
    data["contacts"] = "N/A"
    data["work experience"] = 0 

    try:
        data["work experience"] = openai_extract_experience(\
        getText_docx(data["Name"])
        )
    except Exception as e:
        print("Error in work experience ", data["Name"],  e)



    try:
        text=read_docx_file(data["Name"])
        emails =extract_email_addresses(text)
        if len(emails) > 0 :
            data["email"] = emails[0]

        mobile_numbers = extract_mobile_numbers(text)
        if len(mobile_numbers) > 0 :
            data["contacts"] =  mobile_numbers[0]
    except Exception as e:
        print(e, "\nError processing Email/Mobile Number")
    data["Name"] = os.path.basename(data["Name"])

def getText_docx(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)



def openai_extract_experience(text):
    prompt = "given resume data bellow\n" + text +\
    "\nnwhat is the total number of work experience in years, generate json response, if no work experiance work experiance value should be 0, json format \n{\n\twork experience : < float years>\n}. Output only json data"""
    
    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
                {"role": "system", "content": "You are a chatbot"},
                {"role": "user", "content": prompt},
            ]
    )

    result = ''
    for choice in response.choices:
        result += choice.message.content
    start, end = result.index("{"), result.index("}")
    result = result[start:end+1]
    print(result)
    response = json.loads(result)
    
    return response["work experience"]



